from __future__ import annotations

import math
import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MARKDOWN_PATH = ROOT / "docs" / "软件设计报告-3.4-3.5.md"
DIAGRAM_DIR = ROOT / "docs" / "diagrams"
OUTPUT_PATH = ROOT / "软件设计报告-3.4-3.5.docx"

FONT_PATHS = [
    Path(r"C:\Windows\Fonts\msyh.ttc"),
    Path(r"C:\Windows\Fonts\simhei.ttf"),
    Path(r"C:\Windows\Fonts\simsun.ttc"),
]

INK = "#20303A"
BLUE = "#2E74B5"
DARK_BLUE = "#1F4D78"
LIGHT_BLUE = "#E8EEF5"
LIGHT_GRAY = "#F2F4F7"
MID_GRAY = "#697780"
GREEN = "#2F7D67"
GOLD = "#B8872D"
WHITE = "#FFFFFF"


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path(r"C:\Windows\Fonts\msyhbd.ttc") if bold else Path(r"C:\Windows\Fonts\msyh.ttc"),
        *FONT_PATHS,
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for char in text:
        candidate = current + char
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = char
    if current:
        lines.append(current)
    return lines or [""]


def draw_centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.ImageFont,
    fill: str = INK,
    padding: int = 20,
) -> None:
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, x2 - x1 - padding * 2)
    line_height = font.size + 10
    total_height = line_height * len(lines)
    y = y1 + (y2 - y1 - total_height) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        x = x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height


def draw_round_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    *,
    fill: str = WHITE,
    outline: str = BLUE,
    font: ImageFont.ImageFont,
    radius: int = 16,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=3)
    draw_centered_text(draw, box, text, font)


def draw_diamond(
    draw: ImageDraw.ImageDraw,
    center: tuple[int, int],
    width: int,
    height: int,
    text: str,
    font: ImageFont.ImageFont,
) -> tuple[int, int, int, int]:
    cx, cy = center
    points = [(cx, cy - height // 2), (cx + width // 2, cy), (cx, cy + height // 2), (cx - width // 2, cy)]
    draw.polygon(points, fill="#FFF8E8", outline=GOLD)
    draw.line(points + [points[0]], fill=GOLD, width=3)
    box = (cx - width // 2, cy - height // 2, cx + width // 2, cy + height // 2)
    draw_centered_text(draw, box, text, font, padding=45)
    return box


def draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    label: str | None = None,
    *,
    color: str = MID_GRAY,
    font: ImageFont.ImageFont | None = None,
) -> None:
    draw.line([start, end], fill=color, width=4)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    arrow_len = 18
    wing = 0.55
    p1 = (
        end[0] - arrow_len * math.cos(angle - wing),
        end[1] - arrow_len * math.sin(angle - wing),
    )
    p2 = (
        end[0] - arrow_len * math.cos(angle + wing),
        end[1] - arrow_len * math.sin(angle + wing),
    )
    draw.polygon([end, p1, p2], fill=color)
    if label and font:
        mx = (start[0] + end[0]) / 2
        my = (start[1] + end[1]) / 2
        bbox = draw.textbbox((0, 0), label, font=font)
        draw.rectangle(
            (mx - 8, my - (bbox[3] - bbox[1]) - 5, mx + (bbox[2] - bbox[0]) + 8, my + 5),
            fill=WHITE,
        )
        draw.text((mx, my - (bbox[3] - bbox[1])), label, font=font, fill=color)


def draw_title(draw: ImageDraw.ImageDraw, title: str, subtitle: str, width: int) -> None:
    title_font = get_font(40, bold=True)
    subtitle_font = get_font(22)
    draw.text((70, 45), title, font=title_font, fill=DARK_BLUE)
    draw.text((72, 105), subtitle, font=subtitle_font, fill=MID_GRAY)
    draw.line((70, 150, width - 70, 150), fill=BLUE, width=4)


def build_component_diagram(path: Path) -> None:
    width, height = 1800, 1280
    image = Image.new("RGB", (width, height), WHITE)
    draw = ImageDraw.Draw(image)
    draw_title(draw, "核心构件协作关系", "控制器、业务服务与算法构件分层依赖", width)
    header_font = get_font(26, bold=True)
    class_font = get_font(24, bold=True)
    small_font = get_font(19)

    columns = [
        (90, 300, "控制器层", ["TestController", "ReportController", "RecommendationController", "MatchController"]),
        (650, 300, "业务服务层", ["TestService", "ReportService", "RecommendationService", "MatchService"]),
        (1210, 300, "算法构件层", ["ScoringEngine", "报告构造方法", "RecommendationRanker", "MatchEngine"]),
    ]
    boxes: dict[str, tuple[int, int, int, int]] = {}
    for x, y, title, names in columns:
        draw.rounded_rectangle((x, 205, x + 480, 1140), radius=22, fill="#FAFBFC", outline="#CBD5DB", width=3)
        draw.text((x + 28, 230), title, font=header_font, fill=DARK_BLUE)
        for index, name in enumerate(names):
            top = y + index * 195
            box = (x + 35, top, x + 445, top + 115)
            draw_round_box(draw, box, name, fill=LIGHT_BLUE if index % 2 == 0 else WHITE, font=class_font)
            boxes[name] = box

    pairs = [
        ("TestController", "TestService"),
        ("ReportController", "ReportService"),
        ("RecommendationController", "RecommendationService"),
        ("MatchController", "MatchService"),
        ("TestService", "ScoringEngine"),
        ("ReportService", "报告构造方法"),
        ("RecommendationService", "RecommendationRanker"),
        ("MatchService", "MatchEngine"),
    ]
    for left, right in pairs:
        a, b = boxes[left], boxes[right]
        draw_arrow(draw, (a[2], (a[1] + a[3]) // 2), (b[0], (b[1] + b[3]) // 2))

    draw.text((90, 1190), "实线箭头表示调用或委托关系；算法构件保持无状态，便于独立单元测试。", font=small_font, fill=MID_GRAY)
    image.save(path, quality=95)


def build_recommendation_activity(path: Path) -> None:
    width, height = 1700, 2940
    image = Image.new("RGB", (width, height), WHITE)
    draw = ImageDraw.Draw(image)
    draw_title(draw, "多维度性格—场景标签映射算法活动图", "RecommendationService + RecommendationRanker", width)
    font = get_font(24)
    small = get_font(20)
    box_w, box_h = 760, 105
    cx = width // 2

    def box(y: int, text: str, fill: str = WHITE) -> tuple[int, int, int, int]:
        b = (cx - box_w // 2, y, cx + box_w // 2, y + box_h)
        draw_round_box(draw, b, text, fill=fill, font=font)
        return b

    start = (cx - 210, 190, cx + 210, 270)
    draw_round_box(draw, start, "开始：请求某场景推荐", fill="#EAF6F0", outline=GREEN, font=font, radius=40)
    b1 = box(340, "读取最近一次基础性格测试结果")
    d1 = draw_diamond(draw, (cx, 565), 600, 160, "是否存在画像？", font)
    error = (40, 680, 430, 790)
    draw_round_box(draw, error, "返回：请先完成基础性格测试", fill="#FFF1F0", outline="#C84B4B", font=small)
    error_end = (90, 835, 380, 905)
    draw_round_box(draw, error_end, "结束", fill="#FFF1F0", outline="#C84B4B", font=small, radius=35)
    b2 = box(730, "读取推荐项、用户标签偏好、全局规则权重")
    b3 = box(900, "遍历当前场景的每个推荐项")
    b4 = box(1070, "score = baseScore")
    b5 = box(1240, "遍历推荐项 tags")
    b6 = box(1410, "叠加 userPreference[tag] 与 ruleWeight[tag]")
    d2 = draw_diamond(draw, (cx, 1645), 720, 180, "根据标签映射性格维度", font)

    tag_boxes = [
        (70, 1810, 385, 1935, "explore\n开放性"),
        (465, 1810, 780, 1935, "structured\n尽责性"),
        (860, 1810, 1175, 1935, "social\n外向性"),
        (1255, 1810, 1570, 1935, "gentle\n宜人性"),
    ]
    for x1, y1, x2, y2, text in tag_boxes:
        draw_round_box(draw, (x1, y1, x2, y2), text, fill=LIGHT_BLUE, font=small)

    merge = box(2010, "叠加 (维度分 - 50) / 5；未知标签增加 0")
    d3 = draw_diamond(draw, (cx, 2225), 570, 150, "还有标签？", font)
    clamp = box(2340, "将当前推荐项得分限制在 0～100")
    d4 = draw_diamond(draw, (cx, 2555), 620, 155, "还有推荐项？", font)
    sort_box = box(2670, "按得分降序排序并返回 RecommendationResponse", fill="#EAF6F0")
    end = (cx - 170, 2840, cx + 170, 2915)
    draw_round_box(draw, end, "结束", fill="#EAF6F0", outline=GREEN, font=font, radius=38)

    draw_arrow(draw, (cx, start[3]), (cx, b1[1]))
    draw_arrow(draw, (cx, b1[3]), (cx, d1[1]))
    draw_arrow(draw, (d1[0], (d1[1] + d1[3]) // 2), (error[2], (error[1] + error[3]) // 2), "否", font=small)
    draw_arrow(draw, ((error[0] + error[2]) // 2, error[3]), ((error_end[0] + error_end[2]) // 2, error_end[1]))
    draw_arrow(draw, (cx, d1[3]), (cx, b2[1]), "是", font=small)
    for upper, lower in [(b2, b3), (b3, b4), (b4, b5), (b5, b6)]:
        draw_arrow(draw, (cx, upper[3]), (cx, lower[1]))
    draw_arrow(draw, (cx, b6[3]), (cx, d2[1]))
    for tag_box in tag_boxes:
        tx = (tag_box[0] + tag_box[2]) // 2
        draw_arrow(draw, (cx, d2[3]), (tx, tag_box[1]))
        draw_arrow(draw, (tx, tag_box[3]), (cx, merge[1]))
    draw_arrow(draw, (cx, merge[3]), (cx, d3[1]))
    draw_arrow(draw, (cx, d3[3]), (cx, clamp[1]), "否", font=small)
    draw.line([(d3[2], (d3[1] + d3[3]) // 2), (1600, (d3[1] + d3[3]) // 2), (1600, 1290), (b5[2], 1290)], fill=MID_GRAY, width=4)
    draw_arrow(draw, (1600, 1290), (b5[2], 1290), "是", font=small)
    draw_arrow(draw, (cx, clamp[3]), (cx, d4[1]))
    draw_arrow(draw, (cx, d4[3]), (cx, sort_box[1]), "否", font=small)
    draw.line([(d4[2], (d4[1] + d4[3]) // 2), (1640, (d4[1] + d4[3]) // 2), (1640, 950), (b3[2], 950)], fill=MID_GRAY, width=4)
    draw_arrow(draw, (1640, 950), (b3[2], 950), "是", font=small)
    draw_arrow(draw, (cx, sort_box[3]), (cx, end[1]))
    image.save(path, quality=95)


def build_radar_activity(path: Path) -> None:
    width, height = 1700, 3160
    image = Image.new("RGB", (width, height), WHITE)
    draw = ImageDraw.Draw(image)
    draw_title(draw, "雷达图实时生成计算活动图", "答案提交、分数归一化、报告转换与 ECharts 渲染", width)
    font = get_font(23)
    small = get_font(20)
    cx = width // 2
    box_w, box_h = 800, 100

    def box(y: int, text: str, fill: str = WHITE) -> tuple[int, int, int, int]:
        b = (cx - box_w // 2, y, cx + box_w // 2, y + box_h)
        draw_round_box(draw, b, text, fill=fill, font=font)
        return b

    start = (cx - 210, 185, cx + 210, 265)
    draw_round_box(draw, start, "开始：用户提交答案", fill="#EAF6F0", outline=GREEN, font=font, radius=40)
    b1 = box(330, "TestController 接收 TestSubmitRequest")
    b2 = box(480, "TestService 校验测试类型、题目和选项归属")
    d1 = draw_diamond(draw, (cx, 700), 650, 165, "数据是否合法？", font)
    error = (40, 830, 410, 940)
    draw_round_box(draw, error, "抛出 BusinessException", fill="#FFF1F0", outline="#C84B4B", font=small)
    error_end = (85, 980, 365, 1050)
    draw_round_box(draw, error_end, "结束", fill="#FFF1F0", outline="#C84B4B", font=small, radius=35)
    b3 = box(850, "读取 QuestionOption.weights 并按维度聚合")
    b4 = box(1000, "遍历五个 PersonalityDimension")
    d2 = draw_diamond(draw, (cx, 1225), 700, 170, "当前维度是否有权重？", font)
    default = (90, 1370, 650, 1485)
    average = (1050, 1370, 1610, 1485)
    draw_round_box(draw, default, "否：average = 3.0", fill="#FFF8E8", outline=GOLD, font=font)
    draw_round_box(draw, average, "是：average = 权重平均值", fill=LIGHT_BLUE, font=font)
    b5 = box(1550, "score = round(clamp((average-1)/4×100, 0, 100))")
    d_loop = draw_diamond(draw, (cx, 1765), 570, 150, "还有计分维度？", font)
    b6 = box(1880, "保存 TestResult，ReportService 读取最近结果")
    b_display = box(2030, "遍历报告显示维度")
    d3 = draw_diamond(draw, (cx, 2250), 700, 170, "维度是否为 NEUROTICISM？", font)
    invert = (90, 2395, 720, 2515)
    direct = (980, 2395, 1610, 2515)
    draw_round_box(draw, invert, "是：显示分 = 100 - 原始分\n标签 = 情绪稳定性", fill="#FFF8E8", outline=GOLD, font=small)
    draw_round_box(draw, direct, "否：显示分 = 原始分", fill=LIGHT_BLUE, font=font)
    d_display_loop = draw_diamond(draw, (cx, 2645), 600, 150, "还有显示维度？", font)
    b7 = box(2760, "构造 indicators、radarValues 和解释文本")
    b8 = box(2890, "前端 ECharts setOption() 渲染 radar 并监听 resize", fill="#EAF6F0")
    end = (cx - 170, 3040, cx + 170, 3115)
    draw_round_box(draw, end, "结束", fill="#EAF6F0", outline=GREEN, font=font, radius=38)

    draw_arrow(draw, (cx, start[3]), (cx, b1[1]))
    draw_arrow(draw, (cx, b1[3]), (cx, b2[1]))
    draw_arrow(draw, (cx, b2[3]), (cx, d1[1]))
    draw_arrow(draw, (d1[0], (d1[1] + d1[3]) // 2), (error[2], (error[1] + error[3]) // 2), "否", font=small)
    draw_arrow(draw, ((error[0] + error[2]) // 2, error[3]), ((error_end[0] + error_end[2]) // 2, error_end[1]))
    draw_arrow(draw, (cx, d1[3]), (cx, b3[1]), "是", font=small)
    draw_arrow(draw, (cx, b3[3]), (cx, b4[1]))
    draw_arrow(draw, (cx, b4[3]), (cx, d2[1]))
    draw_arrow(draw, (d2[0], 1225), (default[2], (default[1] + default[3]) // 2), "否", font=small)
    draw_arrow(draw, (d2[2], 1225), (average[0], (average[1] + average[3]) // 2), "是", font=small)
    draw_arrow(draw, ((default[0] + default[2]) // 2, default[3]), (cx, b5[1]))
    draw_arrow(draw, ((average[0] + average[2]) // 2, average[3]), (cx, b5[1]))
    draw_arrow(draw, (cx, b5[3]), (cx, d_loop[1]))
    draw_arrow(draw, (cx, d_loop[3]), (cx, b6[1]), "否", font=small)
    draw.line([(d_loop[2], (d_loop[1] + d_loop[3]) // 2), (1600, (d_loop[1] + d_loop[3]) // 2), (1600, 1050), (b4[2], 1050)], fill=MID_GRAY, width=4)
    draw_arrow(draw, (1600, 1050), (b4[2], 1050), "是", font=small)
    draw_arrow(draw, (cx, b6[3]), (cx, b_display[1]))
    draw_arrow(draw, (cx, b_display[3]), (cx, d3[1]))
    draw_arrow(draw, (d3[0], 1935), (invert[2], (invert[1] + invert[3]) // 2), "是", font=small)
    draw_arrow(draw, (d3[2], 1935), (direct[0], (direct[1] + direct[3]) // 2), "否", font=small)
    draw_arrow(draw, ((invert[0] + invert[2]) // 2, invert[3]), (cx, d_display_loop[1]))
    draw_arrow(draw, ((direct[0] + direct[2]) // 2, direct[3]), (cx, d_display_loop[1]))
    draw_arrow(draw, (cx, d_display_loop[3]), (cx, b7[1]), "否", font=small)
    draw.line([(d_display_loop[2], (d_display_loop[1] + d_display_loop[3]) // 2), (1620, (d_display_loop[1] + d_display_loop[3]) // 2), (1620, 2080), (b_display[2], 2080)], fill=MID_GRAY, width=4)
    draw_arrow(draw, (1620, 2080), (b_display[2], 2080), "是", font=small)
    draw_arrow(draw, (cx, b7[3]), (cx, b8[1]))
    draw_arrow(draw, (cx, b8[3]), (cx, end[1]))
    image.save(path, quality=95)


def set_run_font(run, name: str = "Microsoft YaHei", size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.replace("#", ""))


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "C9D2D8", size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_inches: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_inches:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(round(width * 1440)))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width_dxa = round(widths_inches[index] * 1440)
            cell.width = Inches(widths_inches[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width_dxa))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MID_GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, size=9, color=MID_GRAY)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    heading_specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (11.5, DARK_BLUE, 9, 4),
        "Heading 4": (10.8, DARK_BLUE, 7, 3),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("性格雷达·生活指南  |  软件设计报告")
    set_run_font(run, size=9, color=MID_GRAY)
    p_pr = header._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "5")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "CBD5DB")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = kicker.add_run("软件工程课程设计")
    set_run_font(r, size=12, bold=True, color=GOLD)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(20)
    title.paragraph_format.space_after = Pt(10)
    r = title.add_run("性格雷达·生活指南")
    set_run_font(r, size=28, bold=True, color=DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    r = subtitle.add_run("软件设计报告：3.4 子系统/构件设计与 3.5 详细类设计")
    set_run_font(r, size=15, color=BLUE)

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = line.add_run("基于当前 Vue 3 + Spring Boot 实现")
    set_run_font(r, size=11, color=MID_GRAY)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(140)
    r = note.add_run("核心内容：控制类方法签名、推荐引擎、双人适配、活动图与算法说明")
    set_run_font(r, size=10.5, color=MID_GRAY)
    doc.add_page_break()


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    set_table_borders(table, color="D9E0E5", size="4")
    cell = table.cell(0, 0)
    shade_cell(cell, "F6F8FA")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, name="Consolas", size=8.8, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) > 1 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in rows[1]):
        rows.pop(1)
    return rows


def widths_for_table(headers: list[str]) -> list[float]:
    count = len(headers)
    if count == 4 and headers[:2] == ["可见性", "方法名"]:
        return [0.72, 1.48, 3.15, 1.15]
    if count == 3:
        return [1.45, 2.2, 2.85]
    if count == 2:
        return [2.05, 4.45]
    return [6.5 / count] * count


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    widths = widths_for_table(rows[0])
    set_table_geometry(table, widths)
    set_table_borders(table)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                shade_cell(cell, LIGHT_GRAY)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            run = paragraph.add_run(value.replace("`", ""))
            set_run_font(run, size=8.6 if len(rows[0]) >= 4 else 9, bold=row_index == 0, color=INK)
            if col_index == 0 and len(value) < 18:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_inline_markdown_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run, size=10.5)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=9.2, color=DARK_BLUE)
        else:
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=10.5, bold=True)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=10.5)


def add_diagram(doc: Document, image_path: Path, caption: str, width: float = 6.2) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.keep_with_next = False
    run = cap.add_run(caption)
    set_run_font(run, size=9, color=MID_GRAY)


def markdown_to_docx(doc: Document, text: str, diagram_paths: dict[str, Path]) -> None:
    lines = text.splitlines()
    index = 0
    current_heading = ""
    skip_mermaid = False
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("# "):
            index += 1
            continue
        if stripped.startswith("#### "):
            current_heading = stripped[5:]
            add_inline_markdown_paragraph(doc, current_heading, "Heading 4")
            index += 1
            continue
        if stripped.startswith("### "):
            current_heading = stripped[4:]
            add_inline_markdown_paragraph(doc, current_heading, "Heading 3")
            if current_heading.startswith("3.4.6"):
                add_diagram(doc, diagram_paths["components"], "图 3-8 核心构件协作关系")
                skip_mermaid = True
            elif current_heading.startswith("3.5.10"):
                add_diagram(doc, diagram_paths["recommendation"], "图 3-9 多维度性格—场景标签映射算法活动图", width=5.1)
                skip_mermaid = True
            elif current_heading.startswith("3.5.11"):
                add_diagram(doc, diagram_paths["radar"], "图 3-10 雷达图实时生成计算活动图", width=4.95)
                skip_mermaid = True
            index += 1
            continue
        if stripped.startswith("## "):
            current_heading = stripped[3:]
            add_inline_markdown_paragraph(doc, current_heading, "Heading 1")
            index += 1
            continue
        if stripped.startswith("```"):
            language = stripped[3:].strip()
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            index += 1
            if language == "mermaid" and skip_mermaid:
                skip_mermaid = False
                continue
            add_code_block(doc, code_lines)
            continue
        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_table(doc, parse_table(table_lines))
            continue
        if stripped.startswith("> "):
            table = doc.add_table(rows=1, cols=1)
            set_table_geometry(table, [6.5])
            set_table_borders(table, color="AFC5D8", size="5")
            cell = table.cell(0, 0)
            shade_cell(cell, "F3F7FA")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(stripped[2:])
            set_run_font(run, size=9.6, color=DARK_BLUE)
            index += 1
            continue
        if re.match(r"^\d+\.\s", stripped):
            add_inline_markdown_paragraph(doc, re.sub(r"^\d+\.\s", "", stripped), "List Number")
            index += 1
            continue
        if stripped.startswith("- "):
            add_inline_markdown_paragraph(doc, stripped[2:], "List Bullet")
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            next_line = lines[index].strip()
            if not next_line or next_line.startswith(("#", "|", "```", "> ", "- ")) or re.match(r"^\d+\.\s", next_line):
                break
            paragraph_lines.append(next_line)
            index += 1
        add_inline_markdown_paragraph(doc, " ".join(paragraph_lines))


def build_report() -> None:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    component_path = DIAGRAM_DIR / "core-components.png"
    recommendation_path = DIAGRAM_DIR / "recommendation-activity.png"
    radar_path = DIAGRAM_DIR / "radar-activity.png"
    build_component_diagram(component_path)
    build_recommendation_activity(recommendation_path)
    build_radar_activity(radar_path)

    doc = Document()
    configure_document(doc)
    add_cover(doc)
    markdown_to_docx(
        doc,
        MARKDOWN_PATH.read_text(encoding="utf-8"),
        {
            "components": component_path,
            "recommendation": recommendation_path,
            "radar": radar_path,
        },
    )
    doc.core_properties.title = "性格雷达·生活指南软件设计报告：3.4 与 3.5"
    doc.core_properties.subject = "子系统/构件设计、详细类设计与活动图"
    doc.core_properties.keywords = "软件工程, UML, 活动图, 推荐算法, 雷达图"
    doc.save(OUTPUT_PATH)
    print(f"Generated: {OUTPUT_PATH}")
    print(f"Diagrams: {component_path}, {recommendation_path}, {radar_path}")


if __name__ == "__main__":
    build_report()
