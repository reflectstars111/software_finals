from __future__ import annotations

import shutil
import stat
from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shape import InlineShape
from docx.shared import Cm, Pt
from docxcompose.composer import Composer


ROOT = Path(r"D:\software\finals")
DOC_DIR = ROOT / "设计文档"
TARGET = DOC_DIR / "2024302102007-7485(1).docx"
WORKING = DOC_DIR / "2024302102007-7485(1)-整合中.docx"
BACKUP = DOC_DIR / "2024302102007-7485(1)-整合前备份.docx"
DATA_DESIGN = DOC_DIR / "3.6 数据设计（补全）.docx"
ER_IMAGE = DOC_DIR / "3.6-er-diagram.png"
FRAGMENT_DIR = DOC_DIR / "_merge_fragments"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for style_name, size, east_asia in [
        ("Heading 1", 16, "黑体"),
        ("Heading 2", 14, "黑体"),
        ("Heading 3", 12, "黑体"),
    ]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def make_er_diagram(output: Path) -> None:
    plt.rcParams["font.family"] = "DejaVu Sans"
    fig, ax = plt.subplots(figsize=(14, 9), dpi=180)
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.axis("off")

    nodes = {
        "users": (5.6, 4.1, 2.1, 0.85),
        "question": (0.5, 7.3, 2.0, 0.85),
        "question_option": (3.2, 7.3, 2.4, 0.85),
        "option_weights": (6.3, 7.3, 2.4, 0.85),
        "test_result": (1.1, 4.7, 2.1, 0.85),
        "result_scores": (1.1, 2.8, 2.2, 0.85),
        "report_snapshot": (4.6, 6.1, 2.5, 0.85),
        "share_link": (8.4, 6.6, 2.1, 0.85),
        "recommendation_item": (10.8, 4.8, 2.7, 0.85),
        "feedback": (8.4, 3.6, 2.0, 0.85),
        "user_preference": (10.9, 1.5, 2.5, 0.85),
        "recommendation_rule": (10.8, 7.3, 2.7, 0.85),
        "compatibility_report": (4.1, 1.1, 3.0, 0.85),
        "admin_log": (0.5, 0.8, 2.0, 0.85),
    }

    centers: dict[str, tuple[float, float]] = {}
    for name, (x, y, width, height) in nodes.items():
        centers[name] = (x + width / 2, y + height / 2)
        box = FancyBboxPatch(
            (x, y),
            width,
            height,
            boxstyle="round,pad=0.08,rounding_size=0.08",
            linewidth=1.25,
            edgecolor="#274C5E",
            facecolor="#F4F8F5",
        )
        ax.add_patch(box)
        ax.text(
            x + width / 2,
            y + height / 2,
            name,
            ha="center",
            va="center",
            fontsize=9,
            fontweight="bold",
            color="#16313D",
        )

    relations = [
        ("question", "question_option", "1 : N"),
        ("question_option", "option_weights", "1 : N"),
        ("users", "test_result", "1 : N"),
        ("test_result", "result_scores", "1 : N"),
        ("users", "report_snapshot", "1 : N"),
        ("test_result", "report_snapshot", "1 : N"),
        ("users", "share_link", "1 : N"),
        ("users", "feedback", "1 : N"),
        ("recommendation_item", "feedback", "1 : N"),
        ("users", "user_preference", "1 : N"),
        ("users", "compatibility_report", "owner/target"),
        ("users", "admin_log", "1 : N"),
    ]

    for source, target, label in relations:
        start = centers[source]
        end = centers[target]
        arrow = FancyArrowPatch(
            start,
            end,
            arrowstyle="-|>",
            mutation_scale=10,
            linewidth=1.0,
            color="#718C7A",
            connectionstyle="arc3,rad=0.05",
            shrinkA=22,
            shrinkB=22,
        )
        ax.add_patch(arrow)
        mx = (start[0] + end[0]) / 2
        my = (start[1] + end[1]) / 2
        ax.text(
            mx,
            my,
            label,
            fontsize=7.5,
            ha="center",
            va="center",
            color="#5A675F",
            bbox={"facecolor": "white", "edgecolor": "none", "pad": 0.8},
        )

    ax.set_title(
        "Personality Radar - Core Entity Relationship Diagram",
        fontsize=15,
        fontweight="bold",
        color="#16313D",
        pad=14,
    )
    fig.tight_layout()
    fig.savefig(output, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def add_field_table(
    document: Document,
    table_number: int,
    table_name: str,
    fields: list[tuple[str, str, str]],
) -> None:
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(f"表 3-6-{table_number}  {table_name} 表结构")
    run.bold = True

    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["字段名", "类型", "约束", "说明"]
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "DCE9E2")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    set_repeat_table_header(table.rows[0])

    for name, data_type, constraint, description in fields:
        cells = table.add_row().cells
        values = [name, data_type, constraint, description]
        for index, value in enumerate(values):
            cells[index].text = value
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    document.add_paragraph()


def create_data_design_document() -> None:
    make_er_diagram(ER_IMAGE)
    document = Document()
    style_document(document)

    document.add_heading("3.6 数据设计", level=1)
    document.add_paragraph(
        "本节依据当前 Spring Boot/JPA 领域模型完成 MySQL 物理数据设计。"
        "数据库采用关系模型保存用户、题库、测试结果、报告快照、推荐规则、反馈、分享链接、"
        "双人适配报告和管理员审计日志；Redis 仅用于会话或高频数据缓存，不作为业务事实数据源。"
    )

    document.add_heading("3.6.1 数据设计原则", level=2)
    for text in [
        "一致性：业务实体以 BIGINT 自增主键标识，关联字段使用外键语义，写操作由 Spring 事务控制。",
        "可追溯性：测试、报告、反馈、适配与管理员操作均保留创建时间，不覆盖历史事实记录。",
        "可扩展性：人格维度分数、题目选项权重与推荐标签采用从表存储，便于增加维度或标签。",
        "安全性：密码仅保存 BCrypt 单向散列；账号状态、失败次数与锁定时间支持登录风控。",
        "性能：手机号、分享令牌、推荐规则标签和用户偏好组合键建立唯一约束，常用外键建立索引。",
    ]:
        document.add_paragraph(text, style="List Bullet")

    document.add_heading("3.6.2 概念数据模型与实体关系", level=2)
    document.add_paragraph(
        "核心实体关系如图 3-9 所示。USER 与 TEST_RESULT、REPORT_SNAPSHOT、SHARE_LINK、"
        "FEEDBACK、USER_PREFERENCE、COMPATIBILITY_REPORT 和 ADMIN_LOG 均为一对多关系；"
        "QUESTION 与 QUESTION_OPTION 为一对多关系；选项权重、测试维度分数和推荐标签使用附属表保存。"
    )
    document.add_picture(str(ER_IMAGE), width=Cm(16.0))
    picture_caption = document.add_paragraph("图 3-9 核心实体关系图")
    picture_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("3.6.3 物理表设计", level=2)
    tables: list[tuple[str, list[tuple[str, str, str]]]] = [
        (
            "users",
            [
                ("id", "BIGINT", "PK, AUTO_INCREMENT, NOT NULL", "用户主键"),
                ("phone", "VARCHAR(11)", "UNIQUE, NOT NULL", "登录手机号"),
                ("password_hash", "VARCHAR(255)", "NOT NULL", "BCrypt 密码散列"),
                ("display_name", "VARCHAR(40)", "NOT NULL", "展示名称"),
                ("avatar_url", "VARCHAR(255)", "NULL", "头像地址"),
                ("role", "VARCHAR(20)", "NOT NULL", "USER 或 ADMIN"),
                ("active", "TINYINT(1)", "NOT NULL, DEFAULT 1", "账号是否启用"),
                ("failed_login_attempts", "INT", "NOT NULL, DEFAULT 0", "连续登录失败次数"),
                ("locked_until", "DATETIME(6)", "NULL", "临时锁定截止时间"),
                ("last_login_at", "DATETIME(6)", "NULL", "最近登录时间"),
                ("created_at", "DATETIME(6)", "NOT NULL", "创建时间"),
            ],
        ),
        (
            "question",
            [
                ("id", "BIGINT", "PK, AUTO_INCREMENT", "题目主键"),
                ("type", "VARCHAR(24)", "NOT NULL", "PERSONALITY、FOOD、TRAVEL"),
                ("content", "VARCHAR(240)", "NOT NULL", "题干"),
                ("sort_order", "INT", "NOT NULL", "同类题目排序"),
                ("active", "TINYINT(1)", "NOT NULL", "启用状态"),
            ],
        ),
        (
            "question_option",
            [
                ("id", "BIGINT", "PK, AUTO_INCREMENT", "选项主键"),
                ("question_id", "BIGINT", "FK -> question.id, NOT NULL", "所属题目"),
                ("label", "VARCHAR(20)", "NOT NULL", "选项标识"),
                ("content", "VARCHAR(180)", "NOT NULL", "选项文本"),
                ("sort_order", "INT", "NOT NULL", "选项排序"),
            ],
        ),
        (
            "question_option_weights",
            [
                ("option_id", "BIGINT", "FK -> question_option.id, NOT NULL", "所属选项"),
                ("dimension", "VARCHAR(255)", "PK(联合), NOT NULL", "人格维度代码"),
                ("weight_value", "INT", "NULL", "该选项对维度的权重"),
            ],
        ),
        (
            "test_result",
            [
                ("id", "BIGINT", "PK, AUTO_INCREMENT", "测试结果主键"),
                ("user_id", "BIGINT", "FK -> users.id, NOT NULL", "答题用户"),
                ("type", "VARCHAR(24)", "NOT NULL", "测试类型"),
                ("created_at", "DATETIME(6)", "NOT NULL", "提交时间"),
            ],
        ),
        (
            "test_result_scores",
            [
                ("test_result_id", "BIGINT", "FK -> test_result.id, NOT NULL", "测试结果"),
                ("dimension", "VARCHAR(255)", "PK(联合), NOT NULL", "人格维度代码"),
                ("score", "INT", "NULL", "归一化百分制得分"),
            ],
        ),
        (
            "recommendation_item",
            [
                ("id", "BIGINT", "PK, AUTO_INCREMENT", "推荐项主键"),
                ("scene", "VARCHAR(24)", "NOT NULL", "FOOD、TRAVEL、OUTFIT、CAREER"),
                ("title", "VARCHAR(80)", "NOT NULL", "推荐标题"),
                ("description", "VARCHAR(500)", "NOT NULL", "推荐说明"),
                ("base_score", "INT", "NOT NULL", "推荐基础分"),
                ("active", "TINYINT(1)", "NOT NULL", "启用状态"),
            ],
        ),
        (
            "recommendation_item_tags",
            [
                ("item_id", "BIGINT", "FK -> recommendation_item.id, NOT NULL", "推荐项"),
                ("tag", "VARCHAR(255)", "NULL", "场景标签"),
            ],
        ),
        (
            "recommendation_rule",
            [
                ("id", "BIGINT", "PK, AUTO_INCREMENT", "规则主键"),
                ("tag", "VARCHAR(50)", "UNIQUE, NOT NULL", "标签代码"),
                ("label", "VARCHAR(80)", "NOT NULL", "标签名称"),
                ("weight", "INT", "NOT NULL", "全局规则权重"),
                ("active", "TINYINT(1)", "NOT NULL", "启用状态"),
            ],
        ),
        (
            "user_preference",
            [
                ("id", "BIGINT", "PK, AUTO_INCREMENT", "用户偏好主键"),
                ("user_id", "BIGINT", "FK -> users.id, NOT NULL", "用户"),
                ("tag", "VARCHAR(50)", "UNIQUE(user_id, tag), NOT NULL", "偏好标签"),
                ("weight", "INT", "NOT NULL", "反馈累计权重，范围 -30 到 30"),
            ],
        ),
        (
            "feedback",
            [
                ("id", "BIGINT", "PK, AUTO_INCREMENT", "反馈主键"),
                ("user_id", "BIGINT", "FK -> users.id, NOT NULL", "反馈用户"),
                ("item_id", "BIGINT", "FK -> recommendation_item.id, NOT NULL", "推荐项"),
                ("rating", "VARCHAR(20)", "NOT NULL", "LIKE、DISLIKE 等评价"),
                ("comment", "VARCHAR(100)", "NULL", "补充意见"),
                ("created_at", "DATETIME(6)", "NOT NULL", "反馈时间"),
            ],
        ),
        (
            "report_snapshot",
            [
                ("id", "BIGINT", "PK, AUTO_INCREMENT", "报告快照主键"),
                ("owner_id", "BIGINT", "FK -> users.id, NOT NULL", "报告所有者"),
                ("source_result_id", "BIGINT", "FK -> test_result.id, NOT NULL", "来源测试结果"),
                ("report_json", "TEXT", "NOT NULL", "雷达数据、画像与建议快照"),
                ("summary", "VARCHAR(240)", "NOT NULL", "报告摘要"),
                ("created_at", "DATETIME(6)", "NOT NULL", "生成时间"),
            ],
        ),
        (
            "share_link",
            [
                ("id", "BIGINT", "PK, AUTO_INCREMENT", "分享链接主键"),
                ("owner_id", "BIGINT", "FK -> users.id, NOT NULL", "链接所有者"),
                ("token", "VARCHAR(64)", "UNIQUE, NOT NULL", "不可预测的分享令牌"),
                ("report_json", "TEXT", "NOT NULL", "分享时固化的报告数据"),
                ("active", "TINYINT(1)", "NOT NULL", "链接状态"),
                ("created_at", "DATETIME(6)", "NOT NULL", "创建时间"),
                ("expires_at", "DATETIME(6)", "NULL", "过期时间"),
                ("revoked_at", "DATETIME(6)", "NULL", "撤销时间"),
            ],
        ),
        (
            "compatibility_report",
            [
                ("id", "BIGINT", "PK, AUTO_INCREMENT", "适配报告主键"),
                ("owner_id", "BIGINT", "FK -> users.id, NOT NULL", "发起用户"),
                ("target_id", "BIGINT", "FK -> users.id, NOT NULL", "目标用户"),
                ("score", "DOUBLE", "NOT NULL", "双人契合度"),
                ("summary", "VARCHAR(300)", "NOT NULL", "适配摘要"),
                ("advantages", "VARCHAR(500)", "NOT NULL", "优势说明"),
                ("warnings", "VARCHAR(500)", "NOT NULL", "差异预警"),
                ("advice", "VARCHAR(500)", "NOT NULL", "相处建议"),
                ("created_at", "DATETIME(6)", "NOT NULL", "生成时间"),
            ],
        ),
        (
            "admin_log",
            [
                ("id", "BIGINT", "PK, AUTO_INCREMENT", "日志主键"),
                ("admin_id", "BIGINT", "FK -> users.id, NOT NULL", "操作管理员"),
                ("action", "VARCHAR(80)", "NOT NULL", "操作类型"),
                ("detail", "VARCHAR(500)", "NOT NULL", "操作详情"),
                ("created_at", "DATETIME(6)", "NOT NULL", "操作时间"),
            ],
        ),
    ]

    for table_number, (table_name, fields) in enumerate(tables, start=1):
        add_field_table(document, table_number, table_name, fields)

    document.add_heading("3.6.4 主外键、索引与完整性约束", level=2)
    for text in [
        "所有业务表使用 id 作为主键；附属集合表使用所属实体外键与维度/标签构成逻辑联合键。",
        "users.phone、share_link.token、recommendation_rule.tag 建立唯一约束，防止重复账号、令牌和规则。",
        "user_preference 使用 (user_id, tag) 唯一约束，保证每名用户对每个标签只有一条累计偏好。",
        "question_option.question_id、test_result.user_id、feedback.user_id/item_id 等外键列建立普通索引。",
        "删除题目时由 JPA 级联删除选项及权重；测试、报告、反馈等历史数据原则上不做物理级联删除。",
    ]:
        document.add_paragraph(text, style="List Bullet")

    document.add_heading("3.6.5 敏感数据存储与访问控制", level=2)
    document.add_paragraph(
        "密码字段不采用可逆 AES 加密，而采用 BCrypt 单向散列。注册时由 PasswordEncoder.encode() "
        "生成带随机盐的散列值，登录时通过 matches() 比较，系统无法还原明文密码。"
    )
    document.add_paragraph(
        "手机号用于唯一登录检索，当前实现以受权限控制的 VARCHAR 字段保存。若部署环境要求对手机号、"
        "报告 JSON 等个人信息进行静态加密，可在 Service 与 Repository 之间增加 AES-256-GCM "
        "属性转换器：每条记录使用独立随机 IV，密钥仅由环境变量或密钥管理服务提供，数据库只保存"
        "密文、IV 与认证标签；检索用手机号可额外保存 HMAC-SHA-256 盲索引。"
    )
    document.add_paragraph(
        "分享令牌使用不可预测随机值并设置唯一约束、启用状态、过期时间和撤销时间。管理员接口受"
        " ROLE_ADMIN 权限限制，所有关键管理操作写入 admin_log；数据库账号遵循最小权限原则，"
        "备份文件应加密并限制访问。"
    )

    document.add_heading("3.6.6 数据生命周期与备份", level=2)
    for text in [
        "测试提交后保存 TEST_RESULT 与维度分数；报告生成时写入不可变 REPORT_SNAPSHOT，保证历史报告可复现。",
        "用户反馈写入 FEEDBACK，并同步更新 USER_PREFERENCE；推荐规则调整记录管理员审计日志。",
        "分享链接可主动撤销或按 expires_at 自动失效，失效后保留记录用于审计但不再公开读取。",
        "生产环境每日执行 MySQL 逻辑备份，至少保留 7 个日备份与 4 个周备份，并定期验证恢复流程。",
    ]:
        document.add_paragraph(text, style="List Bullet")

    document.save(DATA_DESIGN)


def element_text(element) -> str:
    return "".join(element.xpath(".//w:t/text()")).strip()


def trim_document(
    source: Path,
    destination: Path,
    start_text: str | None,
    end_text: str | None,
    page_break_before: bool = False,
) -> None:
    shutil.copy2(source, destination)
    destination.chmod(stat.S_IWRITE)
    document = Document(destination)
    body = document._element.body
    children = list(body.iterchildren())

    start_index = 0
    end_index = len(children)
    if start_text is not None:
        for index, child in enumerate(children):
            if child.tag == qn("w:p") and element_text(child) == start_text:
                start_index = index
                break
        else:
            raise RuntimeError(f"Start paragraph not found in {source.name}: {start_text}")
    if end_text is not None:
        for index, child in enumerate(children):
            if child.tag == qn("w:p") and element_text(child) == end_text:
                end_index = index
                break
        else:
            raise RuntimeError(f"End paragraph not found in {source.name}: {end_text}")

    for index, child in reversed(list(enumerate(children))):
        if child.tag == qn("w:sectPr"):
            continue
        if index < start_index or index >= end_index:
            body.remove(child)

    if page_break_before:
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                paragraph.paragraph_format.page_break_before = True
                break
    document.save(destination)


def merge_documents() -> None:
    if not BACKUP.exists():
        shutil.copy2(TARGET, BACKUP)
    FRAGMENT_DIR.mkdir(exist_ok=True)

    prefix = FRAGMENT_DIR / "00-prefix.docx"
    suffix = FRAGMENT_DIR / "99-suffix.docx"
    trim_document(TARGET, prefix, None, "此处插入软件设计说明书…")
    trim_document(TARGET, suffix, "第四部分", None, page_break_before=True)

    fragment_specs = [
        ("01-introduction.docx", DOC_DIR / "设计124(1).docx", "1 引言", "2 软件设计约束和原则"),
        ("02-constraints.docx", DOC_DIR / "设计124(1).docx", "2 软件设计约束和原则", "4 实施指南"),
        (
            "03-architecture.docx",
            DOC_DIR / "3.1 体系结构设计、3.7 部署设计（完整章节）(1).docx",
            "3.1 体系结构设计",
            "3.7 部署设计",
        ),
        ("04-ui.docx", DOC_DIR / "3.2-UI界面设计(2).docx", "3.2 UI 界面设计文档", None),
        ("05-use-cases.docx", DOC_DIR / "D(1).docx", "3.3 用例实现设计", None),
        (
            "06-components-classes.docx",
            DOC_DIR / "软件设计报告-3.4-3.5(1).docx",
            "3.4 子系统/构件设计",
            None,
        ),
        ("07-data.docx", DATA_DESIGN, "3.6 数据设计", None),
        (
            "08-deployment.docx",
            DOC_DIR / "3.1 体系结构设计、3.7 部署设计（完整章节）(1).docx",
            "3.7 部署设计",
            None,
        ),
        ("09-implementation.docx", DOC_DIR / "设计124(1).docx", "4 实施指南", None),
    ]

    fragments: list[Path] = []
    for name, source, start, end in fragment_specs:
        fragment = FRAGMENT_DIR / name
        print(f"PREPARE {name}", flush=True)
        trim_document(source, fragment, start, end, page_break_before=True)
        fragments.append(fragment)

    print("COMPOSE documents", flush=True)
    master = Document(prefix)
    composer = Composer(master)
    for fragment in [*fragments, suffix]:
        print(f"APPEND {fragment.name}", flush=True)
        composer.append(Document(fragment))
    composer.save(WORKING)

    merged = Document(WORKING)
    for index, paragraph in enumerate(merged.paragraphs):
        if paragraph.text.strip() == "3.5.11 雷达图实时生成计算活动图":
            paragraph.paragraph_format.keep_with_next = True
            image_paragraph = merged.paragraphs[index + 1]
            inline_elements = image_paragraph._p.xpath(".//wp:inline")
            if inline_elements:
                shape = InlineShape(inline_elements[0])
                shape.width = int(shape.width * 0.9)
                shape.height = int(shape.height * 0.9)
            break
    merged.save(WORKING)
    WORKING.chmod(stat.S_IWRITE)
    print(f"SAVED {WORKING}", flush=True)


def main() -> None:
    create_data_design_document()
    merge_documents()
    print(f"DATA_DESIGN={DATA_DESIGN}")
    print(f"WORKING={WORKING}")
    print(f"BACKUP={BACKUP}")


if __name__ == "__main__":
    main()
